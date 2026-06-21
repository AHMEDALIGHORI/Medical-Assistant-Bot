"""Generate NLP Lab Report deliverables (sections 9-16) as a Word document."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_shading(hdr_cells[i], "D9E2F3")
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)
    doc.add_paragraph()


def build_document() -> Document:
    doc = Document()

    title = doc.add_heading("NLP Lab Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Medical Assistant Bot — Domain D2 (NLP THEORY CPC)")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True
    doc.add_paragraph()

    # --- Concept ---
    add_heading(doc, "Concept", 1)
    add_para(
        doc,
        "Natural Language Processing (NLP) leverages Python's versatility, simplicity, and extensive "
        "library support to process and analyze human language data. Python's beginner-friendly syntax "
        "and comprehensive NLP libraries make it a preferred choice for professionals and learners "
        "aiming to develop solutions in text analysis, sentiment detection, machine translation, and more.",
    )
    add_para(
        doc,
        "In this project, these capabilities are applied to a Medical Assistant Bot — a web-based "
        "chatbot that ingests medical documents (PDF, CSV, Markdown), preprocesses patient questions, "
        "classifies user intent, retrieves relevant knowledge using Retrieval-Augmented Generation (RAG), "
        "and generates grounded answers through a Streamlit interface powered by Ollama LLMs. Python "
        "ties together every layer: pandas for dataset handling, Hugging Face Transformers for "
        "DistilBERT-based intent classification, sentence-transformers for semantic embeddings, "
        "ChromaDB for vector storage, and Streamlit for deployment.",
    )

    # --- Problem Statement ---
    add_heading(doc, "Problem Statement", 1)
    add_para(
        doc,
        "How can Python be utilized effectively to address complex NLP tasks such as text preprocessing, "
        "named entity recognition, and sentiment analysis? Identify the tools and techniques that "
        "maximize efficiency and accuracy in these tasks.",
    )
    add_para(doc, "Applied to the Medical Assistant Bot, this problem breaks down as follows:", bold=True)
    add_bullet(
        doc,
        "Text preprocessing — Medical text from PDFs, chatbot CSVs, and FAQ files contains OCR noise, "
        "inconsistent whitespace, and unstructured Q/A pairs. Python's regex utilities (clean_medical_text), "
        "word-level chunking (chunk_text), and the DistilBERT subword tokenizer prepare raw text for "
        "both RAG indexing and model training — equivalent to classical tokenization, normalization, "
        "and segmentation tasks performed by NLTK and spaCy.",
    )
    add_bullet(
        doc,
        "Named entity recognition — Patient queries contain medical entities (symptoms, drugs, emergency "
        "indicators). The project uses rule-based keyword detection (infer_intent in ingest_documents.py) "
        "and regex emergency patterns (safety.py) to identify symptoms, medications, and critical conditions. "
        "Hugging Face Transformers further classify the overall intent of each query into structured categories.",
    )
    add_bullet(
        doc,
        "Sentiment / text classification — Rather than positive/negative sentiment, the bot performs "
        "6-class medical intent classification (general_info, symptoms, medication, prevention, emergency, "
        "unknown) using a fine-tuned DistilBERT model. This uses the same supervised classification "
        "paradigm as sentiment analysis, evaluated with accuracy, precision, recall, and F1 score.",
    )
    add_para(
        doc,
        "The project demonstrates that Python, combined with Hugging Face Transformers, scikit-learn, "
        "and sentence-transformers, maximizes both efficiency (modular scripts, fast DistilBERT training "
        "on CPU) and accuracy (96.88% intent classification accuracy on 673 evaluation samples).",
    )

    # --- Design / Ways & Means ---
    add_heading(doc, "Design / Ways & Means", 1)
    add_para(doc, "To explore Python's role in NLP, this lab implements the following in the Medical Assistant Bot:", bold=True)
    add_bullet(
        doc,
        "Employ libraries such as NLTK, spaCy, and Hugging Face Transformers — Hugging Face Transformers "
        "(distilbert-base-uncased) is the primary deep-learning library for intent classification. NLTK/spaCy "
        "techniques (tokenization, stop-word handling, lemmatization) are represented through custom "
        "preprocessing (clean_medical_text, chunk_text) and the built-in DistilBERT WordPiece tokenizer.",
    )
    add_bullet(
        doc,
        "Design workflows for tokenization, stemming, lemmatization, and text classification — "
        "Word-level tokenization via chunk_text (500-word segments, 80-word overlap); subword tokenization "
        "via AutoTokenizer; intent classification via fine-tuned DistilBERT with AdamW optimizer and "
        "stratified 80/20 train-test split.",
    )
    add_bullet(
        doc,
        "Utilize pre-trained models and datasets — Pre-trained distilbert-base-uncased fine-tuned on "
        "3,364 labeled medical questions from ai-medical-chatbot.csv, medical_faq.csv, and symptom datasets; "
        "pre-trained all-MiniLM-L6-v2 for RAG embeddings; Ollama Hermes for response generation.",
    )

    add_heading(doc, "Lab Activities", 2)
    add_bullet(doc, "Activity 1 — Preprocess text data: ingest_documents.py cleans, chunks, and labels medical documents into chunks.jsonl and intent_train.csv.")
    add_bullet(doc, "Activity 2 — Named entity recognition: infer_intent() and safety.py detect medical keywords and emergency entities; DistilBERT classifies query intent.")
    add_bullet(doc, "Activity 3 — Build a sentiment/intent classifier: train_intent_classifier.py fine-tunes DistilBERT and evaluates accuracy (96.88%), precision, recall, and F1.")

    add_heading(doc, "System Architecture", 2)
    add_para(doc, "Pipeline: Ingest → Build Index (ChromaDB) → Train Intent Model → Streamlit Chat App")
    add_table(
        doc,
        ["Component", "Script / Module", "NLP Function"],
        [
            ["Document ingest", "scripts/ingest_documents.py", "Text cleaning, chunking, intent labeling"],
            ["Vector index", "scripts/build_index.py", "Semantic embedding and storage"],
            ["Intent classifier", "scripts/train_intent_classifier.py", "Fine-tune DistilBERT, evaluate metrics"],
            ["Chat interface", "app.py", "RAG retrieval + intent routing + LLM generation"],
            ["Safety layer", "src/safety.py", "Emergency NER via regex patterns"],
        ],
    )

    # --- Analysis & Reporting / Answer ---
    add_heading(doc, "Analysis & Reporting / Answer", 1)
    add_para(
        doc,
        "By conducting lab activities focused on NLP tasks within the Medical Assistant Bot, this report "
        "analyzes Python's efficiency in handling language processing problems. The three lab activities "
        "map directly to the implemented pipeline:",
    )
    add_table(
        doc,
        ["Lab Activity", "Project Implementation", "Key Output"],
        [
            ["Preprocess text (NLTK/spaCy equivalent)", "clean_medical_text + chunk_text + DistilBERT tokenizer", "chunks.jsonl, intent_train.csv (3,364 samples)"],
            ["Named entity recognition (Hugging Face)", "infer_intent keywords + safety.py regex + DistilBERT intent", "6 intent categories, emergency detection"],
            ["Sentiment / intent classifier", "Fine-tuned distilbert-base-uncased", "96.88% accuracy, macro F1 0.7842"],
        ],
    )
    add_para(
        doc,
        "Reporting highlights the strengths of Python's NLP ecosystem in this project: rapid prototyping "
        "with modular scripts, high classification accuracy on medical questions, semantic retrieval over "
        "mixed document formats, and a working Streamlit demo. Areas for improvement include class "
        "imbalance for rare intents (emergency F1 = 0.222), absence of dedicated spaCy/NLTK lemmatization, "
        "and LLM inference latency on CPU. Detailed analysis, results, and discussion follow in Sections 9–16.",
    )

    doc.add_page_break()

    # --- Section 9 ---
    add_heading(doc, "9. Background / Theory", 1)
    add_para(
        doc,
        "Natural Language Processing (NLP) is the branch of artificial intelligence concerned with "
        "enabling computers to understand, interpret, and generate human language. Python has become "
        "the dominant language for NLP because of its readable syntax, strong ecosystem, and integration "
        "with machine learning frameworks.",
    )
    add_para(doc, "In this project, NLP is applied to build a Medical Assistant Bot (Domain D2) that answers health-related questions using:", bold=True)

    add_table(
        doc,
        ["Layer", "Technology", "NLP Role"],
        [
            ["Text cleaning & chunking", "Custom Python (clean_medical_text, chunk_text)", "Normalization, word-level segmentation"],
            ["Subword tokenization", "Hugging Face DistilBERT tokenizer", "Converts text to model-ready tokens"],
            ["Text classification", "Transformers (distilbert-base-uncased)", "Intent routing (symptoms, medication, emergency, etc.)"],
            ["Semantic retrieval", "sentence-transformers + ChromaDB", "Embedding-based document search (RAG)"],
            ["Response generation", "Ollama (Hermes / Qwen)", "Context-conditioned natural language answers"],
        ],
    )

    add_para(doc, "Theoretical foundations:", bold=True)
    add_bullet(doc, "Preprocessing — Raw medical text must be cleaned and split into manageable chunks before indexing or model input, paralleling classical NLTK/spaCy steps.")
    add_bullet(doc, "Named Entity Recognition (NER) — Rule-based intent inference and emergency pattern matching act as lightweight entity/intent detectors in this project.")
    add_bullet(doc, "Text classification — Intent classification into six labels uses the same paradigm as sentiment or topic classification.")
    add_bullet(doc, "Retrieval-Augmented Generation (RAG) — Combines dense retrieval with generative LLMs so answers are grounded in ingested documents.")

    # --- Section 10 ---
    add_heading(doc, "10. Procedure / Methodology", 1)

    add_heading(doc, "Phase 1: Data Ingestion and Preprocessing", 2)
    add_para(doc, "Script: scripts/ingest_documents.py")
    add_para(doc, "Sources scanned: Datafile/ (ai-medical-chatbot.csv, symptom severity, diabetes, heart disease, mental health survey), data/sample/ (medical_faq.csv, general_health.md), and data/raw/ (user uploads).")

    add_table(
        doc,
        ["Step", "Implementation", "Purpose"],
        [
            ["Text cleaning", "clean_medical_text() — removes null bytes, collapses whitespace", "PDF/OCR noise reduction"],
            ["Format normalization", "Q/A pairing: Question: …\\nAnswer: …", "Structured training and retrieval text"],
            ["Intent labeling", "infer_intent() keyword rules or CSV intent column", "Supervised labels for classifier"],
            ["Chunking", "chunk_text() — 500 words, 80-word overlap", "RAG-sized segments"],
            ["Output", "chunks.jsonl, intent_train.csv", "Unified pipeline artifacts"],
        ],
    )

    add_heading(doc, "Phase 2: Vector Index Construction", 2)
    add_bullet(doc, "Script: scripts/build_index.py")
    add_bullet(doc, "Embeddings: all-MiniLM-L6-v2 (sentence-transformers) or Ollama nomic-embed-text")
    add_bullet(doc, "Store: ChromaDB collection medical_knowledge")
    add_bullet(doc, "Retrieval: top-k=4, minimum similarity 0.35")

    add_heading(doc, "Phase 3: Intent Classifier Training (Hugging Face)", 2)
    add_table(
        doc,
        ["Parameter", "Value"],
        [
            ["Base model", "distilbert-base-uncased"],
            ["Max sequence length", "128"],
            ["Train/eval split", "80% / 20% (stratified)"],
            ["Optimizer", "AdamW, LR = 2×10⁻⁵"],
            ["Batch size", "16"],
            ["Epochs", "≥ 5"],
            ["Random seed", "42"],
        ],
    )
    add_para(doc, "Evaluation metrics: accuracy, precision, recall, F1 (macro and per-class via sklearn.classification_report).")

    add_heading(doc, "Phase 4: End-to-End Inference (Streamlit App)", 2)
    add_para(doc, "Script: app.py")
    add_para(doc, "Pipeline: User query → Emergency check (regex) → DistilBERT intent prediction → RAG retrieval → Ollama LLM with intent-specific prompt → Answer + source citations")

    # --- Section 11 ---
    add_heading(doc, "11. ERD (Entity-Relationship Diagram)", 1)
    add_para(doc, "The following describes relationships between text input, preprocessing steps, model training, and output:")
    add_bullet(doc, "One raw document produces many text chunks, which are embedded and stored in ChromaDB.")
    add_bullet(doc, "Chunks with questions generate labeled rows in intent_train.csv for DistilBERT training.")
    add_bullet(doc, "Each user query triggers intent prediction, vector retrieval, and LLM response generation.")
    add_bullet(doc, "Chat output includes the answer, detected intent, confidence, and source citations.")

    add_para(doc, "Entity relationships:", bold=True)
    add_table(
        doc,
        ["Entity", "Key Attributes", "Related To"],
        [
            ["RAW_DOCUMENTS", "source_file, doc_type, raw_text", "PARSED_DOCUMENT"],
            ["TEXT_CHUNK", "chunk_id, text, intent, question", "VECTOR_EMBEDDING, INTENT_TRAINING_ROW"],
            ["DISTILBERT_MODEL", "model_path, macro_f1, num_labels", "TRAINING_SAMPLE"],
            ["USER_QUERY", "query_text, timestamp", "INTENT_PREDICTION, RETRIEVED_CHUNKS"],
            ["CHAT_OUTPUT", "answer, detected_intent, confidence", "LLM_RESPONSE, RETRIEVED_CHUNKS"],
        ],
    )

    # --- Section 12 ---
    add_heading(doc, "12. Analysis", 1)

    add_heading(doc, "Execution Time and Resource Use", 2)
    add_table(
        doc,
        ["Stage", "Typical Behavior", "Notes"],
        [
            ["Document ingest", "Minutes (corpus dependent)", "Fast mode uses 5,000 chatbot rows"],
            ["Index build", "5–15 min (fast) vs 1–2 hrs (full)", "sentence-transformers batch encoding is faster"],
            ["Intent training", "10–30 min on CPU", "DistilBERT is lighter than full BERT"],
            ["Chat inference", "2–5 min first reply on CPU", "Ollama model load dominates latency"],
        ],
    )

    add_heading(doc, "Accuracy and Reliability", 2)
    add_para(doc, "From models/intent_metrics.json (673 evaluation samples):")
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Overall accuracy", "96.88%"],
            ["Macro F1", "0.7842"],
            ["Weighted F1", "0.9643"],
        ],
    )

    add_para(doc, "Per-intent performance:", bold=True)
    add_table(
        doc,
        ["Intent", "Precision", "Recall", "F1", "Support"],
        [
            ["general_info", "0.976", "0.998", "0.987", "493"],
            ["symptoms", "0.938", "0.953", "0.945", "127"],
            ["medication", "0.966", "0.903", "0.933", "31"],
            ["prevention", "1.000", "0.714", "0.833", "14"],
            ["emergency", "1.000", "0.125", "0.222", "8"],
            ["unknown", "—", "—", "0.000", "0"],
        ],
    )

    add_para(
        doc,
        "Interpretation: The model performs strongly on frequent classes (general_info, symptoms) but "
        "under-detects rare classes (emergency, prevention) due to class imbalance. Emergency handling "
        "is supplemented by regex rules in safety.py.",
    )

    add_heading(doc, "Ease of Implementation", 2)
    add_para(doc, "Strengths: unified config.yaml, modular scripts (ingest → index → train → app), standard Hugging Face + scikit-learn metrics, Streamlit sidebar pipeline.")
    add_para(doc, "Trade-offs: custom preprocessing instead of NLTK/spaCy; rule-based labeling introduces noise; RAG quality depends on chunk boundaries and similarity thresholds.")

    # --- Section 13 ---
    add_heading(doc, "13. Results", 1)

    add_heading(doc, "Activity 1: Preprocessed Text", 2)
    add_para(doc, "Input (raw): Q. I have started to get lots of acne on my face, particularly on my forehead. Please help me.")
    add_para(doc, "After cleaning and Q/A formatting: Question and Answer paired with intent label (symptoms).")
    add_para(doc, "After chunking: Text split into ~500-word segments with 80-word overlap, stored in chunks.jsonl.")
    add_para(doc, "Training set: 3,364 labeled questions → 2,691 train / 673 eval.")

    add_heading(doc, "Activity 2: Named Entity / Intent Recognition", 2)
    add_table(
        doc,
        ["User Text Fragment", "Detected Category", "Mechanism"],
        [
            ["chest pain, can't breathe", "emergency", "Keyword list + regex"],
            ["tablet, dose, mg", "medication", "Keyword list"],
            ["fever, cough, rash", "symptoms", "Keyword list"],
            ["vaccine, diet, exercise", "prevention", "Keyword list"],
        ],
    )
    add_para(doc, "Emergency NER (safety.py): Regex patterns for chest pain, stroke, suicide, overdose trigger immediate escalation.")
    add_para(doc, "DistilBERT example: 'What are the side effects of aspirin?' → medication intent with high confidence.")

    add_heading(doc, "Activity 3: Text Classification Results", 2)
    add_para(doc, "Model: Fine-tuned DistilBERT on medical questions. Task: 6-class intent classification.")
    add_bullet(doc, "Accuracy: 96.88%")
    add_bullet(doc, "Best class: general_info (F1 = 0.987)")
    add_bullet(doc, "Weakest class: emergency (F1 = 0.222) — mitigated by regex safety layer")

    add_table(
        doc,
        ["User Question", "Intent", "Outcome"],
        [
            ["What is diabetes?", "general_info", "Educational answer with citations"],
            ["I have fever and cough", "symptoms", "General symptom info + disclaimer"],
            ["chest pain and can't breathe", "emergency (regex)", "Emergency services message"],
        ],
    )

    # --- Section 14 ---
    add_heading(doc, "14. Discussion on Results", 1)

    add_heading(doc, "Strengths", 2)
    add_bullet(doc, "Hugging Face Transformers reached ~97% accuracy with modest compute on 3,000+ medical questions.")
    add_bullet(doc, "sentence-transformers + ChromaDB enabled semantic search over mixed document formats.")
    add_bullet(doc, "Modular pipeline supports easy debugging and extension.")
    add_bullet(doc, "Safety layering combines regex emergency detection, disclaimers, and intent-specific prompts.")
    add_bullet(doc, "Streamlit + Ollama provide a working demo suitable for lab demonstration.")

    add_heading(doc, "Limitations", 2)
    add_bullet(doc, "Class imbalance: rare intents (emergency, prevention) have low recall; macro F1 (0.78) << weighted F1 (0.96).")
    add_bullet(doc, "No classical NLTK/spaCy lemmatization or POS tagging in the preprocessing pipeline.")
    add_bullet(doc, "No dedicated NER model for structured medical entity extraction.")
    add_bullet(doc, "Label noise from infer_intent() auto-labeling propagates into the classifier.")
    add_bullet(doc, "LLM latency on CPU (2–5 min first reply) limits production readiness.")
    add_bullet(doc, "Fixed 500-word chunks may split clinical Q/A pairs awkwardly.")

    add_heading(doc, "Future Improvements", 2)
    add_bullet(doc, "Add spaCy or Hugging Face clinical NER for structured entity extraction.")
    add_bullet(doc, "Oversample or augment rare intent classes.")
    add_bullet(doc, "Use human-verified labels instead of keyword-only labeling.")
    add_bullet(doc, "Evaluate RAG with retrieval precision/recall and answer faithfulness metrics.")

    # --- Section 15 ---
    add_heading(doc, "15. Concluding Remarks", 1)
    add_para(
        doc,
        "This lab demonstrates that Python is well suited for end-to-end NLP in a real application domain — "
        "medical information assistance. The Medical Assistant Bot combines preprocessing, transformer-based "
        "classification, vector retrieval, and generative AI in a single pipeline.",
    )
    add_para(doc, "Key outcomes:", bold=True)
    add_bullet(doc, "3,364 labeled questions processed from mixed medical corpora")
    add_bullet(doc, "96.88% intent classification accuracy on held-out data")
    add_bullet(doc, "Working RAG chatbot with source citations and safety guardrails")
    add_para(
        doc,
        "Python's libraries (Transformers, scikit-learn, sentence-transformers, pandas, Streamlit) allowed "
        "rapid prototyping without sacrificing measurable performance. The project validates the lab objective: "
        "Python can effectively address complex NLP tasks when the right tools and workflow design are applied.",
    )

    # --- Section 16 ---
    add_heading(doc, "16. References", 1)
    refs = [
        "Python Software Foundation. Python Documentation. https://docs.python.org/3/",
        "Bird, S., Klein, E., & Loper, E. (2009). Natural Language Processing with Python (NLTK). O'Reilly Media.",
        "Honnibal, M., & Montani, I. spaCy: Industrial-strength Natural Language Processing. https://spacy.io/",
        "Wolf, T., et al. (2020). Transformers: State-of-the-Art Natural Language Processing. https://huggingface.co/docs/transformers",
        "Sanh, V., et al. (2019). DistilBERT, a distilled version of BERT. https://huggingface.co/distilbert-base-uncased",
        "Reimers, N., & Gurevych, I. (2019). Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks. https://www.sbert.net/",
        "Lewis, P., et al. (2020). Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks. https://arxiv.org/abs/2005.11401",
        "scikit-learn developers. Classification metrics. https://scikit-learn.org/stable/modules/model_evaluation.html",
        "ChromaDB Documentation. https://docs.trychroma.com/",
        "Streamlit Documentation. https://docs.streamlit.io/",
        "Ollama. Local LLM runtime documentation. https://github.com/ollama/ollama",
        "Project repository: medical-assistant-bot/ — Domain D2 Medical Assistant Bot (NLP THEORY CPC).",
    ]
    for i, ref in enumerate(refs, 1):
        doc.add_paragraph(f"{i}. {ref}")

    return doc


def main() -> None:
    out_path = Path(__file__).resolve().parent / "NLP_Lab_Report_Complete.docx"
    doc = build_document()
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
